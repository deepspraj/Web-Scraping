import os

try:
    import requests
    import docx
    from docx2pdf import convert
    from docx.shared import Inches, Cm, Pt
    from bs4 import BeautifulSoup

except ImportError:
    print("Trying to Install required module's \n")
    os.system('pip install requests')
    os.system('pip install docx')
    os.system('pip install docx2pdf')
    os.system('pip install bs4')

class meaningGet:

    def __init__(self, word, path='D:\\vocabulary.docx'):
        self.word = word
        self.path = path
        
    def meaning(self):

        if os.path.isfile(self.path):
            doc = docx.Document(self.path)
        else:
            doc = docx.Document()
            doc.add_heading("Defination's", 0)

        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        url = 'https://www.lexico.com/en/definition/' + self.word
        page = requests.get(url)
        data = page.content

        if (page.status_code == 200):
            soup = BeautifulSoup(data, 'html.parser')
            sect = soup.find_all("section", "gramb")

            for i in range (len(sect)):
                typeOfWord = sect[i].find("span", "pos")
                typeOfWord = typeOfWord.text.title()

                mean = sect[i].find("span", "ind")

                try:    
                    example = sect[i].find("div", "ex")
                    example = example.text
                except:
                    example = 'N.A'

                try:
                    syn1 = sect[i].find("strong","syn")
                    synonyms = syn1.text

                    syn2 = sect[i].find("span", "syn")
                    synonyms = synonyms + syn2.text

                    synonymList = synonyms.split(',')
                    syno = ''
                    for i in range (5):
                        if i != 0:
                            syno = syno + ', ' + synonymList[i]
                        else:
                            syno = synonymList[i]
                except AttributeError:
                    syno = 'N.A'
                
                para = self.word.capitalize() + ' : ' + mean.text + '\n' + 'Example : ' + example + '\n' + 'Synonyms : ' + syno
                
                font = doc.styles['Normal'].font
                font.name = 'Calibri'
                font.size = Pt(14)

                doc.add_paragraph(para)
                doc.save(self.path)
                return 0
        else:
            print('Ooops!! The Website may be under maintenance or shifted to new address (URL).')
            #section is an array with at max 4 element's (adjective, adverb, noun, pronoun) and atleast 1 element(verb).

